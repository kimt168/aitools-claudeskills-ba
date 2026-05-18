"""
DOCX to Markdown Converter
Chuyển đổi file .docx sang định dạng Markdown để AI có thể đọc và xử lý.

Usage:
    python read_docx.py <path_to_docx_file> [--output <output_md_file>]
"""

import sys
import os
import io
import argparse

# Fix encoding for Windows console
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
from typing import Optional
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def is_heading(paragraph) -> tuple[bool, int]:
    """Kiểm tra xem paragraph có phải là heading không và cấp độ của nó."""
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.split()[-1]) if " " in style_name else int(style_name[-1])
            return True, level
        except (ValueError, IndexError):
            return True, 1
    return False, 0


def is_table(element) -> bool:
    """Kiểm tra element có phải là table không."""
    return hasattr(element, 'table')


def extract_list_item(paragraph) -> Optional[str]:
    """Kiểm tra và trích xuất danh sách (bullet/numbered) từ paragraph."""
    if not paragraph._element.xpath('.//w:pPr//w:numPr'):
        return None

    text = paragraph.text.strip()
    if not text:
        return None

    # Xác định loại list (bullet hay numbered)
    numPr = paragraph._element.xpath('.//w:pPr//w:numPr')[0]
    ilvl = numPr.xpath('.//w:ilvl')
    level = int(ilvl[0].get(qn('w:val'))) + 1 if ilvl else 1

    # Xác định prefix dựa trên kiểu list
    numId = numPr.xpath('.//w:numId')
    if numId:
        # Có thể là numbered list - đơn giản hóa dùng bullet
        pass

    indent = "  " * (level - 1)
    return f"{indent}- {text}"


def table_to_markdown(table) -> str:
    """Chuyển đổi bảng Word sang Markdown table."""
    md_lines = []
    rows = []

    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|").replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Tính số cột tối đa
    max_cols = max(len(row) for row in rows)

    # Padding các row để đủ số cột
    for row in rows:
        while len(row) < max_cols:
            row.append("")

    # Header row
    md_lines.append("| " + " | ".join(rows[0]) + " |")

    # Separator
    md_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    # Data rows
    for row in rows[1:]:
        md_lines.append("| " + " | ".join(row) + " |")

    return "\n".join(md_lines)


def docx_to_markdown(docx_path: str) -> str:
    """
    Chuyển đổi file .docx sang Markdown.

    Args:
        docx_path: Đường dẫn đến file .docx

    Returns:
        Chuỗi Markdown chứa nội dung file
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"File không tồn tại: {docx_path}")

    if not docx_path.lower().endswith(".docx"):
        raise ValueError(f"File không phải định dạng .docx: {docx_path}")

    doc = Document(docx_path)
    md_parts = []

    for element in doc.element.body:
        # Xử lý table
        if element.tag.endswith('tbl'):
            table = None
            for tbl in doc.tables:
                if tbl._element is element:
                    table = tbl
                    break
            if table:
                md_table = table_to_markdown(table)
                if md_table:
                    md_parts.append("\n" + md_table + "\n")
            continue

        # Xử lý paragraph
        for para in doc.paragraphs:
            if para._element is element:
                text = para.text.strip()
                if not text:
                    continue

                # Kiểm tra heading
                heading, level = is_heading(para)
                if heading:
                    md_parts.append(f"\n{'#' * level} {text}\n")
                    continue

                # Kiểm tra list item
                list_item = extract_list_item(para)
                if list_item:
                    md_parts.append(list_item + "\n")
                    continue

                # Paragraph thường
                md_parts.append(text + "\n")

    return "\n".join(md_parts)


def extract_metadata(docx_path: str) -> dict:
    """Trích xuất metadata từ file .docx."""
    doc = Document(docx_path)

    paragraphs_count = len([p for p in doc.paragraphs if p.text.strip()])
    tables_count = len(doc.tables)
    headings_count = len([p for p in doc.paragraphs if is_heading(p)[0]])

    return {
        "total_paragraphs": paragraphs_count,
        "total_tables": tables_count,
        "total_headings": headings_count,
        "file_size": os.path.getsize(docx_path),
    }


def convert_file(
    input_path: str,
    output_path: Optional[str] = None,
    include_metadata: bool = True
) -> str:
    """
    Chuyển đổi file .docx sang Markdown và ghi ra file hoặc trả về chuỗi.

    Args:
        input_path: Đường dẫn file .docx đầu vào
        output_path: Đường dẫn file .md đầu ra (nếu None thì trả về chuỗi)
        include_metadata: Có bao gồm metadata vào đầu file không

    Returns:
        Nội dung Markdown
    """
    # Convert
    markdown_content = docx_to_markdown(input_path)

    # Thêm metadata nếu yêu cầu
    if include_metadata:
        metadata = extract_metadata(input_path)
        header = (
            f"---\n"
            f"# Metadata\n"
            f"- Source: {os.path.basename(input_path)}\n"
            f"- Paragraphs: {metadata['total_paragraphs']}\n"
            f"- Tables: {metadata['total_tables']}\n"
            f"- Headings: {metadata['total_headings']}\n"
            f"- File size: {metadata['file_size']} bytes\n"
            f"---\n\n"
        )
        markdown_content = header + markdown_content

    # Ghi ra file hoặc trả về chuỗi
    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        print(f"Đã chuyển đổi thành công: {input_path} -> {output_path}")
    else:
        output_path = os.path.splitext(input_path)[0] + ".md"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        print(f"Đã chuyển đổi thành công: {input_path} -> {output_path}")

    return markdown_content


def main():
    parser = argparse.ArgumentParser(
        description="Chuyển đổi file .docx sang Markdown"
    )
    parser.add_argument(
        "input",
        help="Đường dẫn file .docx đầu vào"
    )
    parser.add_argument(
        "-o", "--output",
        help="Đường dẫn file .md đầu ra (mặc định: cùng tên với file đầu vào)"
    )
    parser.add_argument(
        "--no-metadata",
        action="store_true",
        help="Không bao gồm metadata vào đầu file"
    )

    args = parser.parse_args()

    try:
        convert_file(
            args.input,
            args.output,
            include_metadata=not args.no_metadata
        )
    except FileNotFoundError as e:
        print(f"Lỗi: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"Lỗi: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Lỗi không xác định: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
